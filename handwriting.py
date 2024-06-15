from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)

def text_to_handwritten_image(text, font_path, image_path, font_size=50, max_width=960, max_height=1200):
    
    font = ImageFont.truetype(font_path, font_size)

    image = Image.new('RGB', (max_width, max_height), color='white')
    draw = ImageDraw.Draw(image)

    lines = []

    for paragraph in text.split('\n'):
        words = paragraph.split()
        if not words:
            lines.append('')
        else:
            line = words[0]
            for word in words[1:]:
                # Check if adding the word fits within the width
                if draw.textbbox((0, 0), line + ' ' + word, font=font)[2] <= max_width:
                    line += ' ' + word
                else:
                    lines.append(line)
                    line = word
            lines.append(line)

    y_text = 0
    for line in lines:
        draw.text((0, y_text), line, font=font, fill='black')
        y_text += font_size

    image.save(image_path)

def insert_image_into_doc(image_path, word_doc_path):
    doc = Document()
    
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(6))

    doc.save(word_doc_path)

docx_path = 'C:/Users/SAGAR DEEP/Desktop/Handwriting tool/napoleonI.docx'
font_path = 'C:/Users/SAGAR DEEP/Desktop/Handwriting tool/GreatVibes-Regular.ttf'
image_path = 'C:/Users/SAGAR DEEP/Desktop/Handwriting tool/new_imageV2.jpg'
word_doc_path = 'C:/Users/SAGAR DEEP/Desktop/Handwriting tool/handwritten_generated_doc.docx'

#text = extract_text_from_docx(docx_path)
#text_to_handwritten_image(text, font_path, image_path)

#print(f"handwritten text image saved to {image_path}")

insert_image_into_doc(image_path, word_doc_path)

print(f"Word document with handwritten image saved to {word_doc_path}")